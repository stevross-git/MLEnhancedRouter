"""
RAG (Retrieval-Augmented Generation) System
Provides document ingestion, vector search, and context-aware response generation
"""

import os
import hashlib
import logging
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import json
import uuid

# Document processing
import PyPDF2
import docx
import markdown
from bs4 import BeautifulSoup

# Vector database and embeddings
import chromadb
from chromadb.config import Settings
from chromadb.utils import embedding_functions

# Database models
from sqlalchemy import Column, Integer, String, Text, DateTime, Boolean, Float, JSON
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import relationship

logger = logging.getLogger(__name__)

# Database models will be imported from models.py to avoid circular imports

class RAGSystem:
    """Main RAG system class"""
    
    def __init__(self, db_instance=None, persist_directory="./rag_data"):
        self.db = db_instance
        self.persist_directory = persist_directory
        self.chunk_size = 1000
        self.chunk_overlap = 200
        self.max_chunks_per_query = 5
        
        # Initialize ChromaDB
        self.chroma_client = None
        self.collection = None
        self.embedding_function = None
        
        self._initialize_chroma()
        logger.info("RAG System initialized")
    
    def _initialize_chroma(self):
        """Initialize ChromaDB client and collection"""
        try:
            # Create persist directory if it doesn't exist
            os.makedirs(self.persist_directory, exist_ok=True)
            
            # Initialize ChromaDB client
            self.chroma_client = chromadb.PersistentClient(
                path=self.persist_directory,
                settings=Settings(anonymized_telemetry=False)
            )
            
            # Initialize embedding function (using default sentence transformers)
            self.embedding_function = embedding_functions.DefaultEmbeddingFunction()
            
            # Get or create collection
            try:
                self.collection = self.chroma_client.get_collection(
                    name="documents",
                    embedding_function=self.embedding_function
                )
                logger.info("Loaded existing ChromaDB collection")
            except:
                self.collection = self.chroma_client.create_collection(
                    name="documents",
                    embedding_function=self.embedding_function
                )
                logger.info("Created new ChromaDB collection")
                
        except Exception as e:
            logger.error(f"Failed to initialize ChromaDB: {e}")
            self.chroma_client = None
            self.collection = None
    
    def add_document(self, file_path: str, filename: str, user_id: str = "anonymous", 
                    metadata: Dict[str, Any] = None) -> Optional[str]:
        """Add a document to the RAG system"""
        try:
            # Read file and calculate hash
            with open(file_path, 'rb') as f:
                file_content = f.read()
                file_hash = hashlib.sha256(file_content).hexdigest()
            
            # Import models here to avoid circular imports
            from models import Document, DocumentChunk, RAGQuery
            
            # Check if document already exists
            existing_doc = self.db.session.query(Document).filter_by(file_hash=file_hash).first()
            if existing_doc:
                logger.info(f"Document already exists: {filename}")
                return existing_doc.id
            
            # Determine file type
            file_extension = os.path.splitext(filename)[1].lower()
            file_type = self._get_file_type(file_extension)
            
            # Create document record
            document = Document(
                filename=filename,
                original_name=filename,
                file_type=file_type,
                file_size=len(file_content),
                file_hash=file_hash,
                uploaded_by=user_id,
                metadata=metadata or {}
            )
            
            # Extract text content
            text_content = self._extract_text(file_path, file_type)
            document.content = text_content
            
            # Save document
            self.db.session.add(document)
            self.db.session.commit()
            
            # Process document (create chunks and embeddings)
            self._process_document(document)
            
            logger.info(f"Successfully added document: {filename}")
            return document.id
            
        except Exception as e:
            logger.error(f"Failed to add document {filename}: {e}")
            self.db.session.rollback()
            return None
    
    def _get_file_type(self, extension: str) -> str:
        """Determine file type from extension"""
        type_mapping = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'docx',
            '.txt': 'text',
            '.md': 'markdown',
            '.html': 'html',
            '.htm': 'html',
            '.json': 'json',
            '.csv': 'csv'
        }
        return type_mapping.get(extension, 'unknown')
    
    def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text content from various file types"""
        try:
            if file_type == 'pdf':
                return self._extract_pdf_text(file_path)
            elif file_type == 'docx':
                return self._extract_docx_text(file_path)
            elif file_type == 'text':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            elif file_type == 'markdown':
                with open(file_path, 'r', encoding='utf-8') as f:
                    md_content = f.read()
                    html = markdown.markdown(md_content)
                    return BeautifulSoup(html, 'html.parser').get_text()
            elif file_type == 'html':
                with open(file_path, 'r', encoding='utf-8') as f:
                    html = f.read()
                    return BeautifulSoup(html, 'html.parser').get_text()
            elif file_type == 'json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    return json.dumps(data, indent=2)
            elif file_type == 'csv':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                return ""
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            return ""
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Failed to extract PDF text: {e}")
            return ""
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {e}")
            return ""
    
    def _process_document(self, document):
        """Process document by creating chunks and embeddings"""
        try:
            # Import models here to avoid circular imports
            from models import DocumentChunk
            
            if not self.collection:
                logger.error("ChromaDB collection not initialized")
                return
            
            # Split content into chunks
            chunks = self._split_text(document.content)
            
            # Create document chunks and embeddings
            chunk_ids = []
            chunk_contents = []
            chunk_metadata = []
            
            for i, chunk_content in enumerate(chunks):
                # Create chunk record
                chunk = DocumentChunk(
                    document_id=document.id,
                    chunk_index=i,
                    content=chunk_content,
                    metadata={
                        'document_name': document.original_name,
                        'file_type': document.file_type,
                        'chunk_size': len(chunk_content)
                    }
                )
                
                # Generate unique embedding ID
                embedding_id = f"{document.id}_{i}"
                chunk.embedding_id = embedding_id
                
                self.db.session.add(chunk)
                
                # Prepare for batch embedding
                chunk_ids.append(embedding_id)
                chunk_contents.append(chunk_content)
                chunk_metadata.append({
                    'document_id': document.id,
                    'document_name': document.original_name,
                    'file_type': document.file_type,
                    'chunk_index': i
                })
            
            # Add to ChromaDB collection
            self.collection.add(
                documents=chunk_contents,
                ids=chunk_ids,
                metadatas=chunk_metadata
            )
            
            # Update document status
            document.is_processed = True
            document.processed_at = datetime.utcnow()
            document.chunk_count = len(chunks)
            
            self.db.session.commit()
            
            logger.info(f"Processed document {document.original_name} into {len(chunks)} chunks")
            
        except Exception as e:
            logger.error(f"Failed to process document {document.id}: {e}")
            self.db.session.rollback()
    
    def _split_text(self, text: str) -> List[str]:
        """Split text into chunks with overlap"""
        chunks = []
        words = text.split()
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk = ' '.join(chunk_words)
            chunks.append(chunk)
            
            # Stop if we've reached the end
            if i + self.chunk_size >= len(words):
                break
        
        return chunks
    
    def search_documents(self, query: str, n_results: int = 5, user_id: str = "anonymous") -> List[Dict[str, Any]]:
        """Search documents using vector similarity"""
        try:
            if not self.collection:
                logger.error("ChromaDB collection not initialized")
                return []
            
            # Import models here to avoid circular imports
            from models import RAGQuery
            
            # Generate query hash for caching
            query_hash = hashlib.sha256(query.encode()).hexdigest()
            
            # Check for cached results
            cached_query = self.db.session.query(RAGQuery).filter_by(query_hash=query_hash).first()
            if cached_query and cached_query.retrieved_chunks:
                logger.info(f"Using cached search results for query: {query[:50]}...")
                return cached_query.retrieved_chunks
            
            # Perform vector search
            results = self.collection.query(
                query_texts=[query],
                n_results=min(n_results, self.max_chunks_per_query)
            )
            
            # Process results
            search_results = []
            if results['documents'] and results['documents'][0]:
                for i, (doc, metadata, distance) in enumerate(zip(
                    results['documents'][0],
                    results['metadatas'][0],
                    results['distances'][0]
                )):
                    search_results.append({
                        'content': doc,
                        'metadata': metadata,
                        'similarity_score': 1 - distance,  # Convert distance to similarity
                        'rank': i + 1
                    })
            
            # Save query results
            rag_query = RAGQuery(
                query=query,
                query_hash=query_hash,
                user_id=user_id,
                retrieved_chunks=search_results
            )
            self.db.session.add(rag_query)
            self.db.session.commit()
            
            logger.info(f"Found {len(search_results)} relevant chunks for query: {query[:50]}...")
            return search_results
            
        except Exception as e:
            logger.error(f"Failed to search documents: {e}")
            return []
    
    def generate_context(self, query: str, search_results: List[Dict[str, Any]], 
                        max_context_length: int = 4000) -> str:
        """Generate context string from search results"""
        try:
            if not search_results:
                return ""
            
            context_parts = []
            current_length = 0
            
            # Add search results as context
            for i, result in enumerate(search_results):
                content = result['content']
                metadata = result['metadata']
                
                # Create context chunk with metadata
                context_chunk = f"""
Document: {metadata.get('document_name', 'Unknown')}
Relevance: {result['similarity_score']:.2f}
Content: {content}
---
"""
                
                # Check if adding this chunk would exceed max length
                if current_length + len(context_chunk) > max_context_length:
                    break
                
                context_parts.append(context_chunk)
                current_length += len(context_chunk)
            
            # Combine all context parts
            context = f"""
Based on the following relevant documents, please provide a comprehensive answer to the user's question.

Query: {query}

Relevant Documents:
{''.join(context_parts)}

Please provide a detailed response based on the above context. If the context doesn't contain enough information to answer the question, please mention that and provide what information is available.
"""
            
            return context
            
        except Exception as e:
            logger.error(f"Failed to generate context: {e}")
            return ""
    
    def get_document_stats(self) -> Dict[str, Any]:
        """Get statistics about the document collection"""
        try:
            # Import models here to avoid circular imports
            from models import Document, DocumentChunk, RAGQuery
            from sqlalchemy import func
            
            total_docs = self.db.session.query(Document).count()
            processed_docs = self.db.session.query(Document).filter_by(is_processed=True).count()
            total_chunks = self.db.session.query(DocumentChunk).count()
            total_queries = self.db.session.query(RAGQuery).count()
            
            # Get file type distribution
            file_types = self.db.session.query(
                Document.file_type,
                func.count(Document.id)
            ).group_by(Document.file_type).all()
            
            return {
                'total_documents': total_docs,
                'processed_documents': processed_docs,
                'total_chunks': total_chunks,
                'total_queries': total_queries,
                'file_types': dict(file_types),
                'processing_rate': (processed_docs / total_docs * 100) if total_docs > 0 else 0
            }
            
        except Exception as e:
            logger.error(f"Failed to get document stats: {e}")
            return {}
    
    def delete_document(self, document_id: str) -> bool:
        """Delete a document and its chunks"""
        try:
            # Import models here to avoid circular imports
            from models import Document, DocumentChunk
            
            document = self.db.session.query(Document).get(document_id)
            if not document:
                return False
            
            # Delete from ChromaDB
            if self.collection:
                chunks = self.db.session.query(DocumentChunk).filter_by(document_id=document_id).all()
                embedding_ids = [chunk.embedding_id for chunk in chunks if chunk.embedding_id]
                
                if embedding_ids:
                    self.collection.delete(ids=embedding_ids)
            
            # Delete from database
            self.db.session.query(DocumentChunk).filter_by(document_id=document_id).delete()
            self.db.session.delete(document)
            self.db.session.commit()
            
            logger.info(f"Deleted document {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to delete document {document_id}: {e}")
            self.db.session.rollback()
            return False
    
    def get_documents(self, user_id: str = None, limit: int = 100) -> List[Dict[str, Any]]:
        """Get list of documents"""
        try:
            # Import models here to avoid circular imports
            from models import Document
            
            query = self.db.session.query(Document)
            if user_id:
                query = query.filter_by(uploaded_by=user_id)
            
            documents = query.order_by(Document.uploaded_at.desc()).limit(limit).all()
            
            return [{
                'id': doc.id,
                'filename': doc.original_name,
                'file_type': doc.file_type,
                'file_size': doc.file_size,
                'uploaded_at': doc.uploaded_at.isoformat(),
                'is_processed': doc.is_processed,
                'chunk_count': doc.chunk_count,
                'uploaded_by': doc.uploaded_by
            } for doc in documents]
            
        except Exception as e:
            logger.error(f"Failed to get documents: {e}")
            return []

# Global RAG system instance
rag_system = None

def get_rag_system(db_instance=None) -> RAGSystem:
    """Get or create global RAG system instance"""
    global rag_system
    if rag_system is None:
        rag_system = RAGSystem(db_instance=db_instance)
    return rag_system